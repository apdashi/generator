import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_UNDERLINE
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

bashIPPost = "iptables -t nat -A POSTROUTING -d %s -p tcp -m tcp --dport %s -j SNAT --to-source %s;"
bashIPPre = "iptables -t nat -A PREROUTING -d %s -p tcp -m tcp --dport %s -j DNAT --to-destination %s;"
text1 = "Сайт ОАО «Аэрофлот. План публикации обновления %s. \nНомер заявки-релиза в системе HelpDesk: #%s"
text2_0 = "Выполнить обновление функциональности %s выполненные в рамках заявок:"
text2_1 = "https://support.ramax.ru/issues/%s"
text3 = "Дата и время публикации: с %s до %s %s \n \nСписок контактов ответственных сотрудников:"
text4 = "Особый режим мониторинга не требуется."
text5 = "Системный администратор выполняет резервное копирование приложения и конфигурационных файлов " \
               "на серверах %s "
text6 = "Системный администратор временно перенаправляет трафик %s с сервера %s на %s "
text7 = "Системный администратор производит обновление кода приложения на сервере %s путем переключения" \
        " локальной версии на ветку релиза, выполнив команду Subversion. Номер ревизии: %s "
text8 = "Системный администратор перезапускает приложение на сервере %s и отключает " \
                    "перенаправление трафика "
text9 = "/etc/init.d/prod_schedule restart; \n" \
                "while ! /usr/lib/nagios/plugins/check_tcp -H localhost -p 9780 ; do sleep 0.5 ; done; "
text10 = "Системый администртор осуществляет сброс кеша Varnish на серверах %s"
text12 = "Системный администратор выполняет сравнение предыдущей и новой версии приложения," \
                    " формирует отчет, содержащий следующие сведения: "
text13_0 = "список измененных файлов"
text13_1 = "diff –u (за исключением каталогов logs, pid)"
text13_2 = "исключить из отчета пароли"
text13_3 = "отчет направить на следующие адреса: "
text14 = "Участники релиза совместно осуществляют тестирование по списку тест-кейсов из файла, " \
         "выложенного в релизную заявку. Предварительно тестер у себя прописывает в файл hosts " \
         "(%windir%\System32\drivers\etc\hosts) IP-адрес площадки сайта, и выполняет тесты"
text17 = "%s        www.aeroflot.ru"
text15 = "В случае обнаружения ошибок на этапе тестирования принимается решение об откате" \
                    " изменений на сервере %s. "
text16 = "Производится оповещение ответственных сотрудников (список сформирован на этапе 1) о " \
                "произведенной публикации."
text19_0 = "Производится рассылка плана публикации участникам проекта со стороны Исполнителя (ЗАО «РАМАКС ИНТЕРНЕЙШНЛ»)"
text19_1 = "Команда публикации:"
text20 = "Производится уведомление команды публикации о дате и времени проведения публикации."

"""" формирование документа"""
def saveDoc(self, py):
    self.py = py
    document = docx.Document()
    sections = document.sections
    section = sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.65)
    section.page_width = Inches(8.3)
    section.page_height = Inches(11.7)

    styles = document.styles
    styles["List Number"].font.name = "Calibri"
    styles["Normal"].font.name = "Calibri"

    style = styles.add_style('table_z', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']

    style = styles["table_z"]
    style.font.bold = True

    style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    style.font.name = "Courier New"
    style.font.size = Pt(10)
    styles = document.styles
    style = styles["Code"]

    p = style._element.get_or_add_pPr()
    numPr = OxmlElement('w:shd')  # create number properties element
    numPr.set(qn('w:fill'), "FEF2E8")  # set list type/indentation
    p.append(numPr)

    style = styles.add_style('hyperlink', WD_STYLE_TYPE.CHARACTER)
    style.base_style = styles['Default Paragraph Font']
    style = styles['hyperlink']
    style.font._element.get_or_add_rPr()
    style.font._element.get_or_add_unhideWhenUsed()
    style.font._element.get_or_add_semiHidden()
    style.font._element.rPr.get_or_add_u()
    style.font._element.rPr.u.val = WD_UNDERLINE.SINGLE
    style.font._element.rPr.get_or_add_color()
    style.font._element.rPr.color.val = RGBColor(2, 3, 226)
    style.font._element.rPr.color.themeColor = MSO_THEME_COLOR.HYPERLINK

    h = document.add_header()
    hp = h.add_paragraph(text1 % (self.py.nameProject.itemText(self.py.nameProject.currentIndex()),
                                  self.py.numBid.text()))
    hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(u"Цель", style='table_z')
    document.add_paragraph(text2_0 % (self.py.nameProject.itemText(self.py.nameProject.currentIndex())))
    h = document.add_paragraph('', style='List Bullet 3')
    h.add_run(text2_1 % (self.py.numberBid.text())).style = styles['hyperlink']
    h.add_run(' - %s' % (self.py.specBid.text()))

    document.add_heading("Этап 1. Подготовка к публикации", level=2)
    document.add_paragraph(text3 % (self.py.timeIn.text(), self.py.timeOut.text(), self.py.dateIn.text()))
    make_table(document, self.py.tableView.model().cached, 0)
    document.add_paragraph()
    document.add_paragraph(text19_0)
    document.add_paragraph(text19_1)
    make_table(document, self.py.tableView.model().cached, 1)
    document.add_paragraph()
    document.add_paragraph(text20)
    document.add_paragraph()
    document.add_heading("Этап 2. Публикация изменений", level=2)
    document.add_paragraph()
    listSer = []
    listApp = []
    listVamish = []
    test = []
    if self.py.app1.text() != "":
        listSer = {self.py.app1.text(): self.py.ipAddress1.text()}
        listApp = [self.py.app1.text()]
        listVamish = [self.py.vamish1.text()]
        test = [self.py.test1.text()]
    if self.py.app2.text() != "":
        listSer[self.py.app2.text()] = self.py.ipAddress2.text()
        listApp.append(self.py.app2.text())
        listVamish.append(self.py.vamish2.text())
        test.append(self.py.test2.text())
    if self.py.app3.text() != "":
        listSer[self.py.app3.text()] = self.py.ipAddress3.text()
        listApp.append(self.py.app3.text())
        listVamish.append(self.py.vamish3.text())
        test.append(self.py.test3.text())
    if self.py.app4.text() != "":
        listSer[self.py.app4.text()] = self.py.ipAddress4.text()
        listApp.append(self.py.app4.text())
        listVamish.append(self.py.vamish4.text())
        test.append(self.py.test4.text())
    document.add_paragraph(text5 % ", ".join(i for i in listSer), style='List Number')
    for i in range(0, len(listApp)):
        if (len(listApp) - 1) == i:
            serverPer = 0
        else:
            serverPer = i + 1
        com1 = bashIPPost % (listSer.get(listApp[serverPer]), self.py.portApp.text(), listSer.get(listApp[i]))
        com2 = bashIPPre % (listSer.get(listApp[i]), self.py.portApp.text(), listSer.get(listApp[serverPer]))
        document.add_paragraph(text6 % (self.py.nameProject.itemText(self.py.nameProject.currentIndex()),
                                        listApp[i], listApp[serverPer]), style='List Number')
        document.add_paragraph("%s \n%s" % (com1, com2), style="Code")
        document.add_paragraph(text7 % (listApp[i], self.py.numberRev.text()), style='List Number')
        h = document.add_paragraph('svn switch ', style="Code")
        h.add_run(self.py.svn.text()).style = styles['hyperlink']
        h.add_run('\nsvn update')
        document.add_paragraph(text8 % (listApp[i]), style='List Number')
        document.add_paragraph("%s \n%s \n%s" % (text9, com2, com1), style="Code")
        if self.py.cache.isChecked() == True:
            document.add_paragraph(text10 % (listVamish[i]), style='List Number')
            document.add_paragraph(self.py.comandCache.text(), style="Code")
        document.add_paragraph(text12, style='List Number')
        document.add_paragraph(text13_0, style='List Bullet')
        document.add_paragraph(text13_1, style='List Bullet')
        document.add_paragraph(text13_2, style='List Bullet')
        document.add_paragraph(text13_3, style='List Bullet')
        for j in (self.py.tableView.model().cached):
            if j[2] == True:
                h = document.add_paragraph(style='List Paragraph')
                create_list(h, i)
                h.add_run(j[6]).style = styles['hyperlink']
        document.add_paragraph(text14, style='List Number')
        document.add_paragraph(text17 % (test[i]), style="Code")
        s = ", ".join(c for c in listApp[0:i+1])
        document.add_paragraph(text15 % s, style='List Number')
    document.add_paragraph(text20, style="List Number")
    document.add_heading("Этап 3. Мониторинг работы", level=2)
    document.add_paragraph(text4)

    return document

"""" формирования таблиц в документе"""
def make_table(document, data, num):
    company = []
    for i in data:
        if i[num] == True:
            company.append([i[5], i[4], i[6]])

    company.sort(key=sort_spisok)
    iata = []
    for c in company:
        if c[0].find("RAMAX") != -1 or c[0].find("РАМАКС") != -1:
            iata.insert(0, c)
        else:
            iata.append(c)

    old = "такой компании быть не может"
    tbl = document.add_table(1, 3, style= "Table Grid")
    for i, h in enumerate((u"Компания", u"Представитель", u"Email")):
        tbl.cell(0,i).paragraphs[0].style = document.styles['table_z']
        # tbl.cell(0, i).add_paragraph(h ,style="table_z")
        tbl.cell(0, i).paragraphs[0].text = h
    row_i = 1
    for i, r in enumerate(iata):
        if old != r[0]:
            tbl.add_row()
            tbl.cell(i+row_i, 0).merge(tbl.cell(i+row_i, 2))
            tbl.cell(i+row_i, 0).text = r[0]
            old = r[0]
            row_i += 1
        tbl.add_row()
        for j, val in enumerate(r):
            if j == 2:
                h = tbl.cell(i+row_i, j).paragraphs[0]
                h.add_run(val).style = document.styles['hyperlink']
            if j == 1:
                tbl.cell(i+row_i, j).text = val

def sort_spisok(inputStr):
    return inputStr[0]


def create_list(paragraph, list_type):
    p = paragraph._p #access to xml paragraph element
    pPr = p.get_or_add_pPr() #access paragraph properties
    numPr = OxmlElement('w:numPr') #create number properties element
    numId = OxmlElement('w:ilvl') #create numId element - sets bullet type
    numId.set(qn('w:val'), '1') #set list type/indentation
    numPr.append(numId) #add bullet type to number properties list
    numId = OxmlElement('w:numId') #create numId element - sets bullet type
    numId.set(qn('w:val'), str(list_type+1))
    numPr.append(numId) #add bullet type to number properties list
    pPr.append(numPr) #add number properties to paragraph