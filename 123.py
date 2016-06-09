import docx
from docx.enum.style import WD_STYLE_TYPE

document = docx.Document()
a = document.styles
b = a["Normal"]
styles = document.styles
style= document.styles["List Number"]
document.add_paragraph(" sadas", style="List Number")

document.add_paragraph()
document.add_paragraph(" sadas", style="List Number")

document.add_paragraph("asdsadsad", style="List Number")
document.save("123.docx")